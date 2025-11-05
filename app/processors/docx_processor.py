from pathlib import Path
from typing import Dict, Any, List, Tuple
from docx import Document
from app.utils.logging_config import log
from app.utils.chunking import chunk_text


class DOCXProcessor:
    
    def __init__(self):
        self.supported_extensions = ['.docx']
    
    def can_process(self, file_path: str) -> bool:
        return Path(file_path).suffix.lower() in self.supported_extensions
    
    def extract_text_from_docx(self, docx_path: str) -> str:
        try:
            doc = Document(docx_path)
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            return text.strip()
        except Exception as e:
            log.error(f"Error extracting text from DOCX: {e}")
            return ""
    
    def process(self, file_path: str) -> Tuple[List[str], Dict[str, Any]]:
        try:
            text = self.extract_text_from_docx(file_path)
            
            if not text:
                raise ValueError("No text extracted from DOCX")
            
            chunks_with_indices = chunk_text(text, method="smart")
            chunks = [chunk for chunk, _ in chunks_with_indices]
            
            doc = Document(file_path)
            
            metadata = {
                'file_type': 'docx',
                'file_path': str(file_path),
                'filename': Path(file_path).name,
                'num_paragraphs': len(doc.paragraphs),
                'num_tables': len(doc.tables),
                'text_length': len(text),
                'total_chunks': len(chunks),
                'extension': Path(file_path).suffix
            }
            
            log.info(f"Processed DOCX: {file_path} -> {len(chunks)} chunks")
            return chunks, metadata
            
        except Exception as e:
            log.error(f"Error processing DOCX {file_path}: {e}")
            raise


docx_processor = DOCXProcessor()

__all__ = ["DOCXProcessor", "docx_processor"]